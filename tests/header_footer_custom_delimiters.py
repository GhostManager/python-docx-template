import os
import zipfile

import jinja2
from docx import Document

from docxtpl import DocxTemplate


template_path = "output/header_footer_custom_delimiters_tpl.docx"
output_path = "output/header_footer_custom_delimiters.docx"

os.makedirs("output", exist_ok=True)

document = Document()
document.add_paragraph("Body text")
section = document.sections[0]
section.header.paragraphs[0].text = "[[ date ]]"
section.footer.paragraphs[0].text = "[[ company_name ]]"
document.save(template_path)

jinja_env = jinja2.Environment(
    variable_start_string="[[",
    variable_end_string="]]",
)

tpl = DocxTemplate(template_path)
tpl.render(
    {
        "company_name": "The World Wide company",
        "date": "2016-03-17",
    },
    jinja_env=jinja_env,
)
tpl.save(output_path)

with zipfile.ZipFile(output_path) as docx_zip:
    header_xml = "\n".join(
        docx_zip.read(name).decode("utf-8")
        for name in docx_zip.namelist()
        if name.startswith("word/header")
    )
    footer_xml = "\n".join(
        docx_zip.read(name).decode("utf-8")
        for name in docx_zip.namelist()
        if name.startswith("word/footer")
    )

assert "2016-03-17" in header_xml
assert "The World Wide company" in footer_xml
assert "[[" not in header_xml
assert "[[" not in footer_xml
