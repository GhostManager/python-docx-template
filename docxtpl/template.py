# -*- coding: utf-8 -*-
"""
Created : 2015-03-12

@author: Eric Lapouyade
"""
from __future__ import annotations

from os import PathLike
from typing import TYPE_CHECKING, Any, Optional, IO, Union, Dict, Set
import io
from lxml import etree
from docx import Document
from docx.opc.oxml import parse_xml
from docx.opc.part import XmlPart
import docx.oxml.ns
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as REL_TYPE
from jinja2 import Environment, meta
from jinja2.exceptions import TemplateError


def _create_optimized_env(**kwargs):
    """Create an optimized Jinja2 environment for better performance.
    
    Optimizations applied:
    - auto_reload=False: Skip checking if template source changed
    - cache_size=400: Larger template cache for repeated renders
    - enable_async=False: Disable async support (not needed, adds overhead)
    """
    return Environment(
        auto_reload=False,      # Disable template auto-reload (faster)
        cache_size=400,         # Increase template cache size
        enable_async=False,     # Disable async (not needed, reduces overhead)
        **kwargs
    )


# Module-level cached environments (created once, reused across all instances)
_CACHED_ENV = None
_CACHED_ENV_AUTOESCAPE = None


def _get_cached_env(autoescape=False):
    """Get or create a cached Jinja2 environment for performance."""
    global _CACHED_ENV, _CACHED_ENV_AUTOESCAPE
    
    if autoescape:
        if _CACHED_ENV_AUTOESCAPE is None:
            _CACHED_ENV_AUTOESCAPE = _create_optimized_env(autoescape=True)
        return _CACHED_ENV_AUTOESCAPE
    else:
        if _CACHED_ENV is None:
            _CACHED_ENV = _create_optimized_env(autoescape=False)
        return _CACHED_ENV

try:
    from html import escape  # noqa: F401
except ImportError:
    # cgi.escape is deprecated in python 3.7
    from cgi import escape  # noqa: F401
import re
import binascii
import os
import zipfile

if TYPE_CHECKING:
    from .subdoc import Subdoc


class DocxTemplate(object):
    """Class for managing docx files as they were jinja2 templates"""

    HEADER_URI = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
    )
    FOOTER_URI = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"
    )

    # Pre-compiled regex patterns for patch_xml() optimization
    # These are compiled once at class load time, not on every render
    _RE_JINJA_OPEN = re.compile(
        r"(?<={)(<[^>]*>)+(?=[\{%\#])|(?<=[%\}#])(<[^>]*>)+(?=\})",
        re.DOTALL
    )
    _RE_JINJA_CONTENT = re.compile(
        r"{%(?:(?!%}).)*|{#(?:(?!#}).)*|{{(?:(?!}}).)*",
        re.DOTALL
    )
    _RE_COLSPAN = re.compile(
        r"(<w:tc[ >](?:(?!<w:tc[ >]).)*){%\s*colspan\s+([^%]*)\s*%}(.*?</w:tc>)",
        re.DOTALL
    )
    _RE_CELLBG = re.compile(
        r"(<w:tc[ >](?:(?!<w:tc[ >]).)*){%\s*cellbg\s+([^%]*)\s*%}(.*?</w:tc>)",
        re.DOTALL
    )
    _RE_SPACE_PRESERVE = re.compile(
        r"<w:t>((?:(?!<w:t>).)*)({{.*?}}|{%.*?%})",
        re.DOTALL
    )
    _RE_SPACE_PRESERVE_R = re.compile(
        r"({{r\s.*?}}|{%r\s.*?%})",
        re.DOTALL
    )
    _RE_MERGE_PREV = re.compile(r"</w:t>(?:(?!</w:t>).)*?{%-", re.DOTALL)
    _RE_MERGE_NEXT = re.compile(r"-%}(?:(?!<w:t[ >]|{%|{{).)*?<w:t[^>]*?>", re.DOTALL)
    _RE_VMERGE = re.compile(
        r"<w:tc[ >](?:(?!<w:tc[ >]).)*?{%\s*vm\s*%}.*?</w:tc[ >]",
        re.DOTALL
    )
    _RE_HMERGE = re.compile(
        r"<w:tc[ >](?:(?!<w:tc[ >]).)*?{%\s*hm\s*%}.*?</w:tc[ >]",
        re.DOTALL
    )
    _RE_CLEAN_TAGS = re.compile(r"(?<=\{[\{%])(.*?)(?=[\}%]})")
    _RE_PARAGRAPH_NEWLINE = re.compile(r"<w:p([ >])")
    _RE_PARAGRAPH_REMOVE_NEWLINE = re.compile(r"\n<w:p([ >])")
    _RE_STRIPTAGS = re.compile(r"</w:t>.*?(<w:t>|<w:t [^>]*>)", re.DOTALL)
    _RE_COLSPAN_EMPTY = re.compile(r"<w:r[ >](?:(?!<w:r[ >]).)*<w:t></w:t>.*?</w:r>", re.DOTALL)
    _RE_GRIDSPAN = re.compile(r"<w:gridSpan[^/]*/>")
    _RE_TCPR = re.compile(r"(<w:tcPr[^>]*>)")
    _RE_SHD = re.compile(r"<w:shd[^/]*/>")
    _RE_RESOLVE_PARAGRAPH = re.compile(r"<w:p(?: [^>]*)?>.*?</w:p>", re.DOTALL)
    _RE_RESOLVE_RUN = re.compile(r"<w:r(?: [^>]*)?>.*?</w:r>", re.DOTALL)
    _RE_RESOLVE_TEXT = re.compile(r"<w:t(?: [^>]*)?>.*?</w:t>", re.DOTALL)
    _RE_RUN_PROPS = re.compile(r"<w:rPr>.*?</w:rPr>")
    _RE_PARA_PROPS = re.compile(r"<w:pPr>.*?</w:pPr>")

    # Pre-compiled patterns for tag-stripping in patch_xml().
    # Strips surrounding <w:y> tags from {%y ...%} / {{y ...}} template tags.
    _RE_TAG_STRIP = tuple(
        re.compile(
            r"<w:%s[ >](?:(?!<w:%s[ >]).)*({%%|{{)%s ([^}%%]*(?:%%}|}})).*?</w:%s>"
            % (y, y, y, y),
            re.DOTALL,
        )
        for y in ("tr", "tc", "p", "r")
    )
    # Same for {#y ...#} comment tags (not 'r' - comments in runs are uncommon).
    _RE_COMMENT_STRIP = tuple(
        re.compile(
            r"<w:%s[ >](?:(?!<w:%s[ >]).)*({#)%s ([^}#]*(?:#})).*?</w:%s>"
            % (y, y, y, y),
            re.DOTALL,
        )
        for y in ("tr", "tc", "p")
    )

    # Precompiled pattern for fast detection of any Jinja syntax in a string.
    # Used in render() to skip header/footer processing when no tags are present.
    _JINJA_PATTERN = re.compile(r'\{\{|\{%|\{#')

    def __init__(self, template_file: Union[IO[bytes], str, PathLike]) -> None:
        self.template_file = template_file
        self.reset_replacements()
        self.docx = None
        self.is_rendered = False
        self.is_saved = False
        self.allow_missing_pics = False

    def init_docx(self, reload: bool = True):
        if not self.docx or (self.is_rendered and reload):
            self.docx = Document(self.template_file)
            self.is_rendered = False

    def render_init(self):
        self.init_docx()
        self.pic_map = {}
        self.current_rendering_part = None
        self.docx_ids_index = 1000
        self.is_saved = False

    def __getattr__(self, name):
        return getattr(self.docx, name)

    def xml_to_string(self, xml, encoding="unicode"):
        # Be careful : pretty_print MUST be set to False, otherwise patch_xml()
        # won't work properly
        return etree.tostring(xml, encoding="unicode", pretty_print=False)

    def get_docx(self):
        self.init_docx()
        return self.docx

    def get_xml(self):
        return self.xml_to_string(self.docx._element.body)

    def write_xml(self, filename):
        with open(filename, "w") as fh:
            fh.write(self.get_xml())

    def patch_xml(self, src_xml):
        """Make a lots of cleaning to have a raw xml understandable by jinja2 :
        strip all unnecessary xml tags, manage table cell background color and colspan,
        unescape html entities, etc..."""

        # replace {<something>{ by {{   ( works with {{ }} {% and %} {# and #})
        src_xml = self._RE_JINJA_OPEN.sub("", src_xml)

        # replace {{<some tags>jinja2 stuff<some other tags>}} by {{jinja2 stuff}}
        # same thing with {% ... %} and {# #}
        # "jinja2 stuff" could a variable, a 'if' etc... anything jinja2 will understand
        def striptags(m):
            return self._RE_STRIPTAGS.sub("", m.group(0))

        src_xml = self._RE_JINJA_CONTENT.sub(striptags, src_xml)

        # manage table cell colspan
        def colspan(m):
            cell_xml = m.group(1) + m.group(3)
            cell_xml = self._RE_COLSPAN_EMPTY.sub("", cell_xml)
            cell_xml = self._RE_GRIDSPAN.sub("", cell_xml, count=1)
            return self._RE_TCPR.sub(
                r'\1<w:gridSpan w:val="{{%s}}"/>' % m.group(2),
                cell_xml,
            )

        src_xml = self._RE_COLSPAN.sub(colspan, src_xml)

        # manage table cell background color
        def cellbg(m):
            cell_xml = m.group(1) + m.group(3)
            cell_xml = self._RE_COLSPAN_EMPTY.sub("", cell_xml)
            cell_xml = self._RE_SHD.sub("", cell_xml, count=1)
            return self._RE_TCPR.sub(
                r'\1<w:shd w:val="clear" w:color="auto" w:fill="{{%s}}"/>' % m.group(2),
                cell_xml,
            )

        src_xml = self._RE_CELLBG.sub(cellbg, src_xml)

        # ensure space preservation
        src_xml = self._RE_SPACE_PRESERVE.sub(
            r'<w:t xml:space="preserve">\1\2',
            src_xml,
        )
        src_xml = self._RE_SPACE_PRESERVE_R.sub(
            r'</w:t></w:r><w:r><w:t xml:space="preserve">\1</w:t></w:r><w:r><w:t xml:space="preserve">',
            src_xml,
        )

        # {%- will merge with previous paragraph text
        src_xml = self._RE_MERGE_PREV.sub("{%", src_xml)
        # -%} will merge with next paragraph text
        src_xml = self._RE_MERGE_NEXT.sub("%}", src_xml)

        # Strip surrounding <w:y> tags from {%y ...%} / {{y ...}} template tags.
        # This is mandatory for jinja2 to generate correct xml code.
        # Patterns are pre-compiled as class attributes to avoid recompilation.
        for pat in self._RE_TAG_STRIP:
            src_xml = pat.sub(r"\1 \2", src_xml)

        # Same for {#y ...#} comment tags (not 'r' — comments in runs are uncommon).
        for pat in self._RE_COMMENT_STRIP:
            src_xml = pat.sub(r"\1 \2", src_xml)

        # add vMerge
        # use {% vm %} to make this table cell and its copies
        # be vertically merged within a {% for %}
        def v_merge_tc(m):
            def v_merge(m1):
                return (
                    '<w:vMerge w:val="{% if loop.first %}restart{% else %}continue{% endif %}"/>'
                    + m1.group(1)  # Everything between ``</w:tcPr>`` and ``<w:t>``.
                    + "{% if loop.first %}"
                    + m1.group(2)  # Everything before ``{% vm %}``.
                    + m1.group(3)  # Everything after ``{% vm %}``.
                    + "{% endif %}"
                    + m1.group(4)  # ``</w:t>``.
                )

            return re.sub(
                r"(</w:tcPr[ >].*?<w:t(?:.*?)>)(.*?)(?:{%\s*vm\s*%})(.*?)(</w:t>)",
                v_merge,
                m.group(),
                # Everything between ``</w:tc>`` and ``</w:tc>`` with ``{% vm %}`` inside.
                flags=re.DOTALL,
            )

        src_xml = self._RE_VMERGE.sub(v_merge_tc, src_xml)

        # Use ``{% hm %}`` to make table cell become horizontally merged within
        # a ``{% for %}``.
        def h_merge_tc(m):
            xml_to_patch = (
                m.group()
            )  # Everything between ``</w:tc>`` and ``</w:tc>`` with ``{% hm %}`` inside.

            def with_gridspan(m1):
                return (
                    m1.group(1)  # ``w:gridSpan w:val="``.
                    + "{{ "
                    + m1.group(2)
                    + " * loop.length }}"  # Content of ``w:val``, multiplied by loop length.
                    + m1.group(3)  # Closing quotation mark.
                )

            def without_gridspan(m2):
                return (
                    '<w:gridSpan w:val="{{ loop.length }}"/>'
                    + m2.group(1)  # Everything between ``</w:tcPr>`` and ``<w:t>``.
                    + m2.group(2)  # Everything before ``{% hm %}``.
                    + m2.group(3)  # Everything after ``{% hm %}``.
                    + m2.group(4)  # ``</w:t>``.
                )

            if re.search(r"w:gridSpan", xml_to_patch):
                # Simple case, there's already ``gridSpan``, multiply its value.

                xml = re.sub(
                    r'(w:gridSpan w:val=")(\d+)(")',
                    with_gridspan,
                    xml_to_patch,
                    flags=re.DOTALL,
                )
                xml = re.sub(
                    r"{%\s*hm\s*%}",
                    "",
                    xml,  # Patched xml.
                    flags=re.DOTALL,
                )
            else:
                # There're no ``gridSpan``, add one.
                xml = re.sub(
                    r"(</w:tcPr[ >].*?<w:t(?:.*?)>)(.*?)(?:{%\s*hm\s*%})(.*?)(</w:t>)",
                    without_gridspan,
                    xml_to_patch,
                    flags=re.DOTALL,
                )

            # Discard every other cell generated in loop.
            return "{% if loop.first %}" + xml + "{% endif %}"

        src_xml = self._RE_HMERGE.sub(h_merge_tc, src_xml)

        def clean_tags(m):
            return (
                m.group(0)
                .replace(r"&#8216;", "'")
                .replace("&lt;", "<")
                .replace("&gt;", ">")
                .replace("“", '"')
                .replace("”", '"')
                .replace("‘", "'")
                .replace("’", "'")
            )

        src_xml = self._RE_CLEAN_TAGS.sub(clean_tags, src_xml)

        return src_xml

    def render_xml_part(self, src_xml, part, context, jinja_env=None):
        src_xml = self._RE_PARAGRAPH_NEWLINE.sub(r"\n<w:p\1", src_xml)
        try:
            self.current_rendering_part = part
            if not jinja_env:
                jinja_env = _get_cached_env()
            template = jinja_env.from_string(src_xml)
            dst_xml = template.render(context)
        except TemplateError as exc:
            if hasattr(exc, "lineno") and exc.lineno is not None:
                line_number = max(exc.lineno - 4, 0)
                exc.docx_context = map(
                    lambda x: re.sub(r"<[^>]+>", "", x),
                    src_xml.splitlines()[line_number: (line_number + 7)],  # fmt: skip
                )

            raise exc
        dst_xml = self._RE_PARAGRAPH_REMOVE_NEWLINE.sub(r"<w:p\1", dst_xml)
        dst_xml = (
            dst_xml.replace("{_{", "{{")
            .replace("}_}", "}}")
            .replace("{_%", "{%")
            .replace("%_}", "%}")
        )
        dst_xml = self.resolve_listing(dst_xml)
        return dst_xml

    def render_properties(
        self, context: Dict[str, Any], jinja_env: Optional[Environment] = None
    ) -> None:
        # List of string attributes of docx.opc.coreprops.CoreProperties which are strings.
        # It seems that some attributes cannot be written as strings. Those are commented out.
        properties = [
            "author",
            # 'category',
            "comments",
            # 'content_status',
            "identifier",
            # 'keywords',
            "language",
            # 'last_modified_by',
            "subject",
            "title",
            # 'version',
        ]
        if jinja_env is None:
            jinja_env = _get_cached_env()

        for prop in properties:
            initial = getattr(self.docx.core_properties, prop)
            template = jinja_env.from_string(initial)
            rendered = template.render(context)
            setattr(self.docx.core_properties, prop, rendered)

    def render_footnotes(
        self, context: Dict[str, Any], jinja_env: Optional[Environment] = None
    ) -> None:
        if jinja_env is None:
            jinja_env = _get_cached_env()

        for section in self.docx.sections:
            for part in section.part.package.parts:
                if part.content_type == (
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.footnotes+xml"
                ):
                    xml = self.patch_xml(
                        part.blob.decode("utf-8")
                        if isinstance(part.blob, bytes)
                        else part.blob
                    )
                    xml = self.render_xml_part(xml, part, context, jinja_env)
                    part._blob = xml.encode("utf-8")

    def resolve_listing(self, xml):
        # Early exit: if no Listing special characters are present (common case),
        # there's nothing to resolve, skip the work below.
        if "\t" not in xml and "\n" not in xml and "\a" not in xml and "\f" not in xml:
            return xml

        def resolve_text(run_properties, paragraph_properties, m):
            xml = m.group(0).replace(
                "\t",
                "</w:t></w:r>"
                "<w:r>%s<w:tab/></w:r>"
                '<w:r>%s<w:t xml:space="preserve">' % (run_properties, run_properties),
            )
            xml = xml.replace(
                "\a",
                "</w:t></w:r></w:p>"
                '<w:p>%s<w:r>%s<w:t xml:space="preserve">'
                % (paragraph_properties, run_properties),
            )
            xml = xml.replace("\n", '</w:t><w:br/><w:t xml:space="preserve">')
            xml = xml.replace(
                "\f",
                "</w:t></w:r></w:p>"
                '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'
                '<w:p>%s<w:r>%s<w:t xml:space="preserve">'
                % (paragraph_properties, run_properties),
            )
            return xml

        def resolve_run(paragraph_properties, m):
            run_properties = self._RE_RUN_PROPS.search(m.group(0))
            run_properties = run_properties.group(0) if run_properties else ""
            return self._RE_RESOLVE_TEXT.sub(
                lambda x: resolve_text(run_properties, paragraph_properties, x),
                m.group(0),
            )

        def resolve_paragraph(m):
            paragraph_properties = self._RE_PARA_PROPS.search(m.group(0))
            paragraph_properties = (
                paragraph_properties.group(0) if paragraph_properties else ""
            )
            return self._RE_RESOLVE_RUN.sub(
                lambda x: resolve_run(paragraph_properties, x),
                m.group(0),
            )

        xml = self._RE_RESOLVE_PARAGRAPH.sub(resolve_paragraph, xml)

        return xml

    def build_xml(self, context, jinja_env=None):
        xml = self.get_xml()
        xml = self.patch_xml(xml)
        xml = self.render_xml_part(xml, self.docx._part, context, jinja_env)
        return xml

    def map_tree(self, tree):
        """Replace the body element with the rendered tree.

        Instead of iterating over all body children to remove/re-append them
        one-by-one (O(n) lxml operations, each with internal bookkeeping),
        we swap the entire <w:body> element in the document root using
        root.remove() + root.insert(). This is O(1) since the root element
        (<w:document>) has only a handful of direct children.

        The body's index is located first so document element order is
        preserved (e.g. body before sectPr).

        SAFETY: If the body is not a direct child of root (malformed template)
        or if remove/insert raises for any reason, we fall back to the slower
        child-by-child copy so rendering is never broken.
        """
        root = self.docx._element
        old_body = root.body

        # Find where <w:body> sits among root's direct children so we can
        # re-insert the new tree at the same position.
        body_index = None
        for i, child in enumerate(root):
            if child is old_body:
                body_index = i
                break

        if body_index is None:
            # Malformed template – body is not a direct child of root.
            # Fall back to child-by-child replacement on the existing body.
            for child in list(old_body):
                old_body.remove(child)
            for child in list(tree):
                old_body.append(child)
            return

        try:
            # Detach the old body and insert the new tree (which is itself a
            # <w:body> element returned by fix_tables/parse_xml) at the same
            # position. This avoids O(n) per-child remove/append calls.
            root.remove(old_body)
            root.insert(body_index, tree)
        except Exception:
            # If something went wrong, restore the document to a usable state
            # by re-attaching the old body (if it was already detached) and
            # falling back to child-by-child copy.
            if old_body.getparent() is None:
                root.insert(body_index, old_body)
            for child in list(old_body):
                old_body.remove(child)
            for child in list(tree):
                old_body.append(child)

    def get_headers_footers(self, uri):
        for relKey, val in self.docx._part.rels.items():
            if (val.reltype == uri) and (val.target_part.blob):
                yield relKey, val.target_part

    def get_part_xml(self, part):
        return self.xml_to_string(parse_xml(part.blob))

    def get_headers_footers_encoding(self, xml):
        m = re.match(r'<\?xml[^\?]+\bencoding="([^"]+)"', xml, re.I)
        if m:
            return m.group(1)
        return "utf-8"

    def build_headers_footers_xml(self, context, uri, jinja_env=None):
        for relKey, part in self.get_headers_footers(uri):
            xml = self.get_part_xml(part)
            
            # Skip rendering if no Jinja tags present
            # Headers/footers are often static, so this avoids caching/parsing overhead
            if self._RE_JINJA_OPEN.search(xml) or self._RE_JINJA_CONTENT.search(xml):
                encoding = self.get_headers_footers_encoding(xml)
                xml = self.patch_xml(xml)
                xml = self.render_xml_part(xml, part, context, jinja_env)
                yield relKey, xml.encode(encoding)
            else:
                encoding = self.get_headers_footers_encoding(xml)
                yield relKey, xml.encode(encoding)

    def map_headers_footers_xml(self, relKey, xml):
        part = self.docx._part.rels[relKey].target_part
        new_part = XmlPart.load(part.partname, part.content_type, xml, part.package)
        for rId, rel in part.rels.items():
            new_part.load_rel(rel.reltype, rel._target, rel.rId, rel.is_external)
        self.docx._part.rels[relKey]._target = new_part

    def render(
        self,
        context: Dict[str, Any],
        jinja_env: Optional[Environment] = None,
        autoescape: bool = False,
    ) -> None:
        # init template working attributes
        self.render_init()

        # Use cached environment by default
        if not jinja_env:
            jinja_env = _get_cached_env(autoescape=autoescape)
        elif autoescape:
            jinja_env.autoescape = autoescape

        # Body
        xml_src = self.build_xml(context, jinja_env)

        # Fix tables if needed
        tree = self.fix_tables(xml_src)

        # Fix docPr ID's
        self.fix_docpr_ids(tree)

        # Replace body xml tree
        self.map_tree(tree)

        # Headers & Footers - skip when no Jinja tags are present.
        # Uses both _JINJA_PATTERN (intact tags) and _RE_JINJA_OPEN (tags
        # split across XML runs by Word).
        for uri in (self.HEADER_URI, self.FOOTER_URI):
            try:
                has_jinja = any(
                    self._JINJA_PATTERN.search(xml)
                    or self._RE_JINJA_OPEN.search(xml)
                    for xml in (
                        self.get_part_xml(part)
                        for _relKey, part in self.get_headers_footers(uri)
                    )
                )
                if has_jinja:
                    for relKey, xml in self.build_headers_footers_xml(context, uri, jinja_env):
                        self.map_headers_footers_xml(relKey, xml)
            except Exception:
                # Fallback: guards against unexpected part structure (e.g. blob
                # is None, missing attributes). Not malformed XML - that would
                # fail in build_headers_footers_xml regardless.
                for relKey, xml in self.build_headers_footers_xml(context, uri, jinja_env):
                    self.map_headers_footers_xml(relKey, xml)

        # Properties: no skip-check needed - these are a handful of short
        # strings (author, title, etc.) where from_string() is near-zero cost.
        self.render_properties(context, jinja_env)

        # Footnotes: no skip-check needed - at most one part exists in typical
        # documents, and many have none, so the loop body rarely executes.
        self.render_footnotes(context, jinja_env)

        # set rendered flag
        self.is_rendered = True

    # Using of TC tag in for cycle can cause that count of columns does not
    # correspond to real count of columns in row.
    def fix_tables(self, xml):
        # Use parse_xml with safe fallback for malformed XML
        try:
            tree = parse_xml(xml) # parse_xml() is significantly faster
        except Exception:
            # Fallback to permissive parser in the event of malformed XML
            parser = etree.XMLParser(recover=True)
            tree = etree.fromstring(xml, parser=parser)
        # get namespace
        ns = "{" + tree.nsmap["w"] + "}"
        # walk trough xml and find table
        for t in tree.iter(ns + "tbl"):
            tblGrid = t.find(ns + "tblGrid")
            if tblGrid is None:
                continue
                
            columns = tblGrid.findall(ns + "gridCol")
            columns_len = len(columns)
            
            # Single pass row analysis with both counters
            max_raw_cells = 0       # For ADD decision (raw tc count)
            max_effective_cells = 0  # For REMOVE decision (with gridSpan)
            
            for r in t.iter(ns + "tr"):
                cells = r.findall(ns + "tc")
                raw_count = len(cells)
                effective_count = 0
                
                for cell in cells:
                    tc_pr = cell.find(ns + "tcPr")
                    if tc_pr is not None:
                        grid_span = tc_pr.find(ns + "gridSpan")
                        if grid_span is not None:
                            effective_count += int(grid_span.get(ns + "val"))
                            continue
                    effective_count += 1
                
                if raw_count > max_raw_cells:
                    max_raw_cells = raw_count
                if effective_count > max_effective_cells:
                    max_effective_cells = effective_count
            
            # ADD columns based on RAW cell count (original behavior)
            to_add = max_raw_cells - columns_len if max_raw_cells > columns_len else 0
            
            # is necessary to add columns?
            if to_add > 0:
                # at first, calculate width of table according to columns
                # (we want to preserve it)
                width = 0.0
                new_average = None
                for c in columns:
                    if c.get(ns + "w") is not None:
                        width += float(c.get(ns + "w"))
                # try to keep proportion of table
                if width > 0:
                    old_average = width / len(columns)
                    new_average = width / (len(columns) + to_add)
                    # scale the old columns
                    for c in columns:
                        c.set(
                            ns + "w",
                            str(
                                int(float(c.get(ns + "w")) * new_average / old_average)
                            ),
                        )
                    # add new columns using OxmlElement for proper python-docx compatibility
                    for i in range(to_add):
                        new_col = OxmlElement('w:gridCol')
                        new_col.set(qn('w:w'), str(int(new_average)))
                        tblGrid.append(new_col)

            # REMOVE columns based on EFFECTIVE cell count (original behavior)
            columns = tblGrid.findall(ns + "gridCol")
            columns_len = len(columns)
            to_remove = columns_len - max_effective_cells if columns_len > max_effective_cells else 0

            # If after the loop, there're less columns, than
            # originally was, remove extra `gridCol` declarations.
            if to_remove > 0:
                # Have to keep track of the removed width to scale the
                # table back to its original width.
                removed_width = 0.0

                for c in columns[-to_remove:]:
                    removed_width += float(c.get(ns + "w"))

                    tblGrid.remove(c)

                columns_left = tblGrid.findall(ns + "gridCol")

                # Distribute `removed_width` across all columns that has
                # left after extras removal.
                extra_space = 0
                if len(columns_left) > 0:
                    extra_space = removed_width / len(columns_left)
                    extra_space = int(extra_space)

                for c in columns_left:
                    c.set(ns + "w", str(int(float(c.get(ns + "w")) + extra_space)))

        return tree

    def fix_docpr_ids(self, tree):
        # Some Ids may have some collisions : so renumbering all of them
        wp_ns = docx.oxml.ns.nsmap['wp']
        tag = "{%s}docPr" % wp_ns
        
        for elt in tree.iter(tag):
            self.docx_ids_index += 1
            elt.attrib["id"] = str(self.docx_ids_index)

    def new_subdoc(self, docpath=None) -> Subdoc:
        from .subdoc import Subdoc

        self.init_docx()
        return Subdoc(self, docpath)

    @staticmethod
    def get_file_crc(file_obj):
        if hasattr(file_obj, "read"):
            buf = file_obj.read()
        else:
            with open(file_obj, "rb") as fh:
                buf = fh.read()

        crc = binascii.crc32(buf) & 0xFFFFFFFF
        return crc

    def replace_media(self, src_file, dst_file):
        """Replace one media by another one into a docx

        This has been done mainly because it is not possible to add images in
        docx header/footer.
        With this function, put a dummy picture in your header/footer,
        then specify it with its replacement in this function using the file path
        or file-like objects.

        Syntax: tpl.replace_media('dummy_media_to_replace.png','media_to_paste.jpg')
            -- or --
                tpl.replace_media(io.BytesIO(image_stream), io.BytesIO(new_image_stream))

        Note: for images, the aspect ratio will be the same as the replaced image

        Note2: it is important to have the source media file as it is required
                to calculate its CRC to find them in the docx
        """

        crc = self.get_file_crc(src_file)
        if hasattr(dst_file, "read"):
            self.crc_to_new_media[crc] = dst_file.read()
        else:
            with open(dst_file, "rb") as fh:
                self.crc_to_new_media[crc] = fh.read()

    def replace_pic(self, embedded_file, dst_file):
        """Replace embedded picture with original-name given by embedded_file.
           (give only the file basename, not the full path)
           The new picture is given by dst_file (either a filename or a file-like
           object)

        Notes:
            1) embedded_file and dst_file must have the same extension/format
               in case dst_file is a file-like object, no check is done on
               format compatibility
            2) the aspect ratio will be the same as the replaced image
            3) There is no need to keep the original file (this is not the case
               for replace_embedded and replace_media)
        """

        if hasattr(dst_file, "read"):
            # NOTE: file extension not checked
            self.pics_to_replace[embedded_file] = dst_file.read()
        else:
            with open(dst_file, "rb") as fh:
                self.pics_to_replace[embedded_file] = fh.read()

    def replace_embedded(self, src_file, dst_file):
        """Replace one embedded object by another one into a docx

        This has been done mainly because it is not possible to add images
        in docx header/footer.
        With this function, put a dummy picture in your header/footer,
        then specify it with its replacement in this function

        Syntax: tpl.replace_embedded('dummy_doc.docx','doc_to_paste.docx')

        Note2 : it is important to have the source file as it is required to
                calculate its CRC to find them in the docx
        """
        with open(dst_file, "rb") as fh:
            crc = self.get_file_crc(src_file)
            self.crc_to_new_embedded[crc] = fh.read()

    def replace_zipname(self, zipname, dst_file):
        """Replace one file in the docx file

        First note that a MSWord .docx file is in fact a zip file.

        This method can be used to replace document embedded in the docx template.

        Some embedded document may have been modified by MSWord while saving
        the template : thus replace_embedded() cannot be used as CRC is not the
        same as the original file.

        This method works for embedded MSWord file like Excel or PowerPoint file,
        but won't work for others like PDF, Python or even Text files :
        For these ones, MSWord generate an oleObjectNNN.bin file which is no
        use to be replaced as it is encoded.

        Syntax:

        tpl.replace_zipname(
            'word/embeddings/Feuille_Microsoft_Office_Excel1.xlsx',
            'my_excel_file.xlsx')

        The zipname is the one you can find when you open docx with WinZip,
        7zip (Windows) or unzip -l (Linux). The zipname starts with
        "word/embeddings/". Note that the file is renamed by MSWord,
        so you have to guess a little bit...
        """
        with open(dst_file, "rb") as fh:
            self.zipname_to_replace[zipname] = fh.read()

    def reset_replacements(self):
        """Reset replacement dictionaries

        This will reset data for image/embedded/zipname replacement

        This is useful when calling several times render() with different
        image/embedded/zipname replacements without re-instantiating
        DocxTemplate object.
        In this case, the right sequence for each rendering will be :
            - reset_replacements(...)
            - replace_zipname(...), replace_media(...) and/or replace_embedded(...),
            - render(...)

        If you instantiate DocxTemplate object before each render(),
        this method is useless.
        """
        self.crc_to_new_media = {}
        self.crc_to_new_embedded = {}
        self.zipname_to_replace = {}
        self.pics_to_replace = {}

    def post_processing(self, docx_file):
        if self.crc_to_new_media or self.crc_to_new_embedded or self.zipname_to_replace:

            if hasattr(docx_file, "read"):
                tmp_file = io.BytesIO()
                DocxTemplate(docx_file).save(tmp_file)
                tmp_file.seek(0)
                docx_file.seek(0)
                docx_file.truncate()
                docx_file.seek(0)

            else:
                tmp_file = "%s_docxtpl_before_replace_medias" % docx_file
                os.rename(docx_file, tmp_file)

            with zipfile.ZipFile(tmp_file) as zin:
                with zipfile.ZipFile(docx_file, "w") as zout:
                    for item in zin.infolist():
                        buf = zin.read(item.filename)
                        if item.filename in self.zipname_to_replace:
                            zout.writestr(item, self.zipname_to_replace[item.filename])
                        elif (
                            item.filename.startswith("word/media/")
                            and item.CRC in self.crc_to_new_media
                        ):
                            zout.writestr(item, self.crc_to_new_media[item.CRC])
                        elif (
                            item.filename.startswith("word/embeddings/")
                            and item.CRC in self.crc_to_new_embedded
                        ):
                            zout.writestr(item, self.crc_to_new_embedded[item.CRC])
                        else:
                            zout.writestr(item, buf)

            if not hasattr(tmp_file, "read"):
                os.remove(tmp_file)
            if hasattr(docx_file, "read"):
                docx_file.seek(0)

    def pre_processing(self):

        if self.pics_to_replace:
            self._replace_pics()

    def _replace_pics(self):
        """Replaces pictures xml tags in the docx template with pictures provided by the user"""

        replaced_pics = {key: False for key in self.pics_to_replace}

        # Main document
        part = self.docx.part
        self._replace_docx_part_pics(part, replaced_pics)

        # Header/Footer
        for relid, rel in part.rels.items():
            if rel.reltype in (REL_TYPE.HEADER, REL_TYPE.FOOTER):
                self._replace_docx_part_pics(rel.target_part, replaced_pics)

        if not self.allow_missing_pics:
            # make sure all template images defined by user were replaced
            for img_id, replaced in replaced_pics.items():
                if not replaced:
                    raise ValueError(
                        "Picture %s not found in the docx template" % img_id
                    )

    def get_pic_map(self):
        return self.pic_map

    def _replace_docx_part_pics(self, doc_part, replaced_pics):

        et = etree.fromstring(doc_part.blob)

        part_map = {}

        gds = et.xpath("//a:graphic/a:graphicData", namespaces=docx.oxml.ns.nsmap)
        for gd in gds:
            rel = None
            # Either IMAGE, CHART, SMART_ART, ...
            try:
                if gd.attrib["uri"] == docx.oxml.ns.nsmap["pic"]:
                    # Either PICTURE or LINKED_PICTURE image
                    blip = gd.xpath(
                        "pic:pic/pic:blipFill/a:blip", namespaces=docx.oxml.ns.nsmap
                    )[0]
                    dest = blip.xpath("@r:embed", namespaces=docx.oxml.ns.nsmap)
                    if len(dest) > 0:
                        rel = dest[0]
                    else:
                        continue
                else:
                    continue

                non_visual_properties = "pic:pic/pic:nvPicPr/pic:cNvPr/"
                filename = gd.xpath(
                    "%s@name" % non_visual_properties, namespaces=docx.oxml.ns.nsmap
                )[0]
                titles = gd.xpath(
                    "%s@title" % non_visual_properties, namespaces=docx.oxml.ns.nsmap
                )
                if titles:
                    title = titles[0]
                else:
                    title = ""
                descriptions = gd.xpath(
                    "%s@descr" % non_visual_properties, namespaces=docx.oxml.ns.nsmap
                )
                if descriptions:
                    description = descriptions[0]
                else:
                    description = ""

                part_map[filename] = (
                    doc_part.rels[rel].target_ref,
                    doc_part.rels[rel].target_part,
                )

                # replace data
                for img_id, img_data in self.pics_to_replace.items():
                    if img_id == filename or img_id == title or img_id == description:
                        part_map[filename][1]._blob = img_data
                        replaced_pics[img_id] = True
                        break

            # FIXME: figure out what exceptions are thrown here
            # and catch more specific exceptions
            except Exception:
                continue

        self.pic_map.update(part_map)

    def build_url_id(self, url):
        self.init_docx()
        return self.docx._part.relate_to(url, REL_TYPE.HYPERLINK, is_external=True)

    def save(self, filename: Union[IO[bytes], str, PathLike], *args, **kwargs) -> None:
        # case where save() is called without doing rendering
        # ( user wants only to replace image/embedded/zipname )
        if not self.is_saved and not self.is_rendered:
            self.docx = Document(self.template_file)
        self.pre_processing()
        self.docx.save(filename, *args, **kwargs)
        self.post_processing(filename)
        self.is_saved = True

    def get_undeclared_template_variables(
        self,
        jinja_env: Optional[Environment] = None,
        context: Optional[Dict[str, Any]] = None,
    ) -> Set[str]:
        # Create a temporary document to analyze the template without affecting the current state
        temp_doc = Document(self.template_file)

        # Get XML from the temporary document
        xml = self.xml_to_string(temp_doc._element.body)
        xml = self.patch_xml(xml)

        # Add headers and footers
        for uri in [self.HEADER_URI, self.FOOTER_URI]:
            for relKey, val in temp_doc._part.rels.items():
                if (val.reltype == uri) and (val.target_part.blob):
                    _xml = self.xml_to_string(parse_xml(val.target_part.blob))
                    xml += self.patch_xml(_xml)

        if jinja_env:
            env = jinja_env
        else:
            env = _get_cached_env()

        parse_content = env.parse(xml)
        all_variables = meta.find_undeclared_variables(parse_content)

        # If context is provided, return only variables that are not in the context
        if context is not None:
            provided_variables = set(context.keys())
            return all_variables - provided_variables

        # If no context provided, return all variables (original behavior)
        return all_variables
